from docx import Document
import sys

place, time, *names = sys.stdin

document = Document()

for name in names:
    document.add_heading(f'Добрый день, {name.strip()}!', 1)
    document.add_paragraph(f'Приглашаю тебя на мероприятие в честь 8 марта, которое состоится {place.strip().lower()} {time.strip().lower()}.')
    document.add_page_break()

document.save('res.docx')
