# пришлось переписывать программу с нуля. Тем не менее..
from docx import Document
import sys


def get_lines(text: str) -> list[str]:
    return list(map(lambda line: line.strip(), text.split('\n')))[1:]


def parse(text: str) -> tuple:
    if text.startswith('#'):
        header_level: int = 1
        for i in range(5):
            next_char: str = text[i + 1]
            if next_char == '#':
                header_level += 1
        return text[header_level + 1:], {'type': 'header', 'level': header_level}
    elif len(text) == 0:
        return text, {'type': 'blank'}
    elif text[:text.find('.')].isdigit():
        return text[text.find('.') + 2:], {'type': 'list', 'list-type': 'List Number'}
    elif text.startswith(('+ ', '- ', '* ')):
        return text[2:], {'type': 'list', 'list-type': 'List Bullet'}
    return text, {'type': None}


def apply_style(text: str, paragraph) -> None:
    is_style = False
    while True:
        pos = text.find('*')
        level = text[pos:pos + 3].count('*')
        part = paragraph.add_run(text[:pos] if pos != -1 else text)
        if level == 1:
            part.italic = True
        elif level == 2:
            part.bold = True
        elif level == 3:
            part.italic = True
            part.bold = True
        else:
            part
        text = text[pos + level:]
        is_style = not(is_style)
        if pos == -1:
            break


def markdown_to_docx(text: str) -> None:
    file_name = text[:text.find('\n')]
    lines: list[str] = get_lines(text)
    parsed_lines: list[str|dict] = list(map(lambda line: parse(line), lines))

    doc = Document()
    
    for text, attrs in parsed_lines:
        if attrs['type'] == 'header':
            doc.add_heading(text, level=attrs['level'])
        elif attrs['type'] == 'list':
            p = doc.add_paragraph(style=attrs['list-type'])
            apply_style(text, p)
        else:
            p = doc.add_paragraph()
            apply_style(text, p)

    doc.save(f'{file_name}.docx')


markdown_to_docx(sys.stdin.read())
