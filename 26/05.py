# подло не давать примеры ввода для такой задачи

import sys
from docx import Document


def get_style(text: str) -> str:
    level = 0
    for i in range(len(text) // 2 - 1):
        if text[i] == text[-i - 1] == '*':
            level += 1
    if level == 0:
        return (text, 'normal')
    if level == 1:
        return (text[1:-1], 'italic')
    if level == 2:
        return (text[2:-2], 'bold')
    return (text[3:-3], 'both')


def do_style(paragraph, text, style):
    runner = paragraph.add_run(text)
    if style == 'bold' or style == 'both':
        runner.bold = True
    if style == 'italic' or style == 'both':
        runner.italic = True        


def markdown_to_docx(text: str):
    lines = list(map(lambda line: line.strip(), text))
    marked_lines = []
    for line in lines:
        if not line: 
            marked_lines.append((line, (None, None)))
        elif line.startswith('#'):
            header_level = 1
            for i in range(1, len(line)):
                if line[i] == line[i - 1] == '#':
                    header_level += 1
            marked_lines.append((line[header_level + 1:], ('header', {'level': header_level})))
        elif line[0].isdigit() and line[1] == '.':
            new_line, style = get_style(line[3:])
            marked_lines.append((new_line, ('list', {'type': 'List Number', 'style': style})))
        elif line.startswith(('-', '+')):
            new_line, style = get_style(line[2:])
            marked_lines.append((new_line, ('list', {'type': 'List Bullet', 'style': style})))
        else:
            new_line, style = get_style(line)
            marked_lines.append((new_line, ('paragraph', {'style': style})))

    doc = Document()
    for line, (type, extra) in marked_lines:
        if type == 'header':
            doc.add_heading(line, level=extra['level'])
        elif type == 'list':
            paragraph = doc.add_paragraph(style=extra['type'])
            do_style(paragraph, line, extra['style'])
        elif type == 'paragraph':
            paragraph = doc.add_paragraph()
            do_style(paragraph, line, extra['style'])
        else:
            doc.add_paragraph(line)

    doc.save('res.docx')


markdown_to_docx(sys.stdin.readlines())
